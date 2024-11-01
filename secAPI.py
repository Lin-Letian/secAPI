from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from requests_html import HTMLSession as Session
import os, re, sys, random, platform, hashlib
from typing import Union, List, Tuple, Dict
from urllib.parse import unquote, urlparse
from urllib3 import disable_warnings
from concurrent import futures as cf
from requests import put, get
from docx.shared import RGBColor
from datetime import datetime
from base64 import b64encode
from json import load, dump
from loguru import logger
from docx import Document
from time import time
import base64
from log import log

if not os.path.exists('config/config.json'):
    with open('config/config.json', 'w') as f:
        dump({"token": "", "vInfo": {}}, f)
try:
    with open('config/config.json', 'r') as f:
        CONF = load(f)
except:
    raise '未发现config/config.json文件，请创建并确保存在token'

Timeout = 60
disable_warnings()
Model = {
    "auth": "统一/认证系统",
    "cms": "CMS",
    "Components": "组件",
    "Database": "数据库",
    "Framework": "框架",
    "NetDrivers": "网络设备",
    "OA": "办公系统",
    "ServicesApplication": "服务器应用",
    "SecurityApp": "安全产品",
    "VideoPanel": "视频监控平台",
    "WebApplication": "Web应用程序/系统",
    "WebServer": "Web中间件/服务器"
}
level = {0: "无危害", 1: '低危', 2: '中危', 3: '高危'}
p_poc, p_doc = 'poc', 'output_vul'
API = 'http://127.0.0.1:8080'
disable_warnings()

TOKEN = CONF['token']


def get_header():  # 获取请求头
    ip = f"101.{random.randint(1, 255)}.{random.randint(1, 255)}.{random.randint(1, 255)}"
    _ua = [
        'Mozilla/5.0 (Windows NT 6.1) AppleWebKit/536.3 (HTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3',
        'Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (HTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3',
        'Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.6 (HTML, like Gecko) Chrome/20.0.1090.0 Safari/536.6',
        'Mozilla/5.0 (Windows NT 5.1) AppleWebKit/536.3 (HTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3',
        'Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (HTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3',
        'Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (HTML, like Gecko) Chrome/19.0.1061.0 Safari/536.3',
    ]
    header = {
        'User-Agent': random.choice(_ua),
        "CLIENT-IP": ip,
        "X-FORWARDED-FOR": ip,
        "Accept-Encoding": "gzip, deflate",
        "Accept-Language": "zh-CN,zh;q=0.9",
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9'
    }
    return header


def format_url(url: str):
    tmp = urlparse(url)
    url = tmp.scheme + "://" + tmp.netloc + unquote(tmp.path).replace('../', '.%2E/')
    if tmp.query: url += "?" + tmp.query
    if tmp.fragment: url += "#" + tmp.fragment
    return url


def g_headers(url: str, header: dict = None):
    _header = get_header()
    try:
        ufo = urlparse(url)
        _header.update({"Host": ufo.netloc})
        _header.update({"Referer": ufo.scheme + '://' + ufo.netloc})
    except ValueError as err:
        print(f"Error parsing URL: {url}, Error: {err}")
    if header: _header.update(header)
    return _header


def post(url: str, data=None, json: dict = None, files=None, timeout: int = Timeout, allow_redirects: bool = False,
         verify: bool = False, headers: dict = None, params=None, stream: bool = True, *args):
    r"""
    Sends a POST request.
    :return: :class:`Response <Response>` object
    :rtype: requests.Response
    """
    url = format_url(url)
    if headers is not None and 'Content-Type' not in headers and json is None and files is None:
        headers.update({"Content-Type": "application/x-www-form-urlencoded"})
    headers = g_headers(url=url, header=headers)
    try:
        with Session() as S:
            return S.request(method='post', url=url, data=data, json=json, files=files, allow_redirects=allow_redirects,
                             timeout=timeout,
                             verify=verify, headers=headers, params=params, stream=stream, *args)
    except:
        return None


class Fun:
    @classmethod
    def get_ip(cls):
        try:
            return get('https://cz88.net/api/cz88/ip/base?ip').json()['data']['ip']
        except:
            return '0.0.0.0'

    @classmethod
    def login(cls):
        global TOKEN
        log.info('[*] 登陆')
        user = input('请输入用户名:')
        pas = input('请输入密码:')
        ip = cls.get_ip()
        log.info(f'当前IP: {ip}')
        data = {"Ip": ip, 'uName': user, "Passwd": pas}
        try:
            r = post(API + "/auth/login", json=data)
            TOKEN = r.json()['token']
            with open('config/config.json', 'r') as _f:
                _c = load(_f)
            _c['token'] = TOKEN
            with open('config/config.json', 'w') as _f:
                dump(_c, _f)
            log.info(f'获取新Token: {TOKEN}')
        except Exception as err:
            TOKEN = ''
            log.warning(f'登陆失败，ERR:: {err}')
        return TOKEN

    @classmethod
    def md5(cls, input_string: str) -> str:
        md5_hash = hashlib.md5()
        md5_hash.update(input_string.encode('utf-8'))
        return md5_hash.hexdigest()

    @classmethod
    def b64(cls, text: str) -> str:
        return b64encode(text.encode('utf-8')).decode()

    @classmethod
    def get_computer_hash(cls):
        # 获取系统信息
        system_info = {
            'system': platform.system(),
            'release': platform.release(),
            'machine': platform.machine(),
            'processor': platform.processor(),
        }
        # 将信息转换为字符串并计算哈希值
        info_str = str(system_info)
        hash_object = hashlib.sha256(info_str.encode())
        hex_dig = hash_object.hexdigest()
        return hex_dig

    @classmethod
    def timer(cls) -> str:  # 时间
        now = datetime.now()
        year = f'{now.year:02}'
        month = f'{now.month:02}'
        day = f'{now.day:02}'
        hour = f'{now.hour:02}'
        minute = f'{now.minute:02}'
        second = f'{now.second:02}'
        return f"{year}年{month}月{day}日{hour}时{minute}分{second}秒"

    @classmethod
    def _w_icp(cls, doc, text: str) -> dict:
        dom = urlparse(text).netloc.split(':')[0]
        if dom[0] in ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9']: return doc
        res = post(API + "/api/query/icp", data={'keyword': dom}, headers={'token': TOKEN})
        if res is None: return doc
        try:
            icp = res.json()['data']
        except:
            return doc
        if icp['size'] > 0:
            doc.add_paragraph()
            doc.add_paragraph()
            heading = doc.add_heading('ICP备案查询信息', 1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            icp = icp['list'][0]
            cls._row_tit(doc, "单位名称").add_run(icp['unitName'])
            cls._row_tit(doc, "单位性质").add_run(icp['natureName'])
            cls._row_tit(doc, "ICP号").add_run(icp['serviceLicence'])
            cls._row_tit(doc, "域名").add_run(icp['domain'])
            cls._row_tit(doc, "更新时间").add_run(icp['updateRecordTime'])
        return doc

    @classmethod
    def _row_tit(cls, doc, text: str):
        line = doc.add_paragraph()
        tit = line.add_run(f'{text}：')
        tit.bold = True
        tit.font.color.rgb = RGBColor(0, 51, 102)  # 使用RGB颜色
        return line

    @classmethod
    def write_docx(cls, v: dict, url: str, o_file: str) -> str:
        doc = Document()
        # 标题
        heading = doc.add_heading('漏洞报告', 0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        t = doc.add_paragraph('生成时间：' + Fun.timer())
        t.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        cls._row_tit(doc, "漏洞名称").add_run(v['vName'])
        cls._row_tit(doc, "漏洞URL").add_run(url)
        cls._row_tit(doc, "是否存在漏洞").add_run('是' if v['request'] else '否')
        cls._row_tit(doc, "影响产品").add_run(v['product'])
        cls._row_tit(doc, "影响版本").add_run(v['version'])
        cls._row_tit(doc, "危险等级").add_run(v['level'])
        if v['vId']:
            cls._row_tit(doc, "漏洞编号")
            for n, i in enumerate(v['vId']):
                doc.add_paragraph(str(n + 1) + '：' + i)

        cls._row_tit(doc, "漏洞描述").add_run(v['vDesc'])
        cls._row_tit(doc, "参考链接").add_run(v['link'])
        if v['fix']:
            cls._row_tit(doc, "修复建议")
            for n, i in enumerate(v['fix']):
                doc.add_paragraph(str(n + 1) + '：' + i)
        doc.add_paragraph('')
        doc.add_paragraph('')
        request = v['request']
        cls._row_tit(doc, ">> 请求")
        cls._row_tit(doc, "URL").add_run(request['url'])
        cls._row_tit(doc, "请求方法").add_run(request['method'])
        cls._row_tit(doc, "请求头")
        table = doc.add_table(rows=1, cols=1)
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                hed, text = request['header'], ''
                for i in hed: text += i + ': ' + hed[i] + '\n'
                cell.text = text.rstrip('\n')
        if request['body']:
            cls._row_tit(doc, "请求体")
            table = doc.add_table(rows=1, cols=1)
            for row_index, row in enumerate(table.rows):
                for col_index, cell in enumerate(row.cells): cell.text = request['body']
        doc.add_paragraph('')
        doc.add_paragraph('')
        response = v['response']
        cls._row_tit(doc, ">> 响应")
        cls._row_tit(doc, "URL").add_run(response['url'])
        cls._row_tit(doc, "状态码").add_run(str(response['code']))
        cls._row_tit(doc, "响应头")
        table = doc.add_table(rows=1, cols=1)
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                hed, text = response['header'], ''
                for i in hed: text += i + ': ' + hed[i] + '\n'
                cell.text = text.rstrip('\n')
        cls._row_tit(doc, "响应体")
        table = doc.add_table(rows=1, cols=1)
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells): cell.text = response['body']
        cls._w_icp(doc, url)
        doc.save(os.path.join(p_doc, o_file, v['vName'].replace("/", '') + '.docx'))
        return v['vName'] + '.docx'


class secAPI:
    @classmethod
    def login(cls, ):
        global TOKEN
        log.info('[*] 登陆')
        user = input('请输入用户名:')
        pas = input('请输入密码:')
        ip = cls.self_ip()
        log.info(f'当前IP: {ip}')
        data = {"Ip": ip, 'uName': user, "Passwd": pas}
        try:
            r = post(API + "/auth/login", json=data)
            TOKEN = r.json()['token']
            with open('config/config.json', 'r') as _f:
                C = load(_f)
            C['token'] = TOKEN
            with open('config/config.json', 'w') as _f:
                dump(C, _f)
            log.info(f'获取新Token: {TOKEN}')
        except Exception as err:
            TOKEN = ''
            log.warning(f'登陆失败，ERR:: {err}')
        return TOKEN

    # 获取当前公网IP
    @classmethod
    def self_ip(cls, ) -> str:
        try:
            return get('https://cz88.net/api/cz88/ip/base?ip').json()['data']['ip']
        except:
            return '0.0.0.0'

    # 返回请求用到的header头
    @classmethod
    def g_headers(cls, ) -> dict:
        return {'token': TOKEN, "Content-Type": "application/x-www-form-urlencoded"}

    # 获取备案信息
    @classmethod
    def get_beian(cls, domain: str) -> Union[dict, list]:
        global TOKEN
        data = "keyword={}".format(domain)
        res = post(API + "/api/query/icp", data=data, headers=cls.g_headers())
        if res is None: return {}
        if res.json()['code'] == 401:
            TOKEN = cls.login()
            return cls.get_beian(domain)
        try:
            return res.json()['data']  # 假设接口返回的是JSON格式的数据
        except:
            return {}

    # 获取url解析的IP
    @classmethod
    def site_get_ip(cls, url: str) -> str:
        global TOKEN
        data = "type=ip&site={}".format(base64.b64encode(url.encode('utf-8')).decode())
        res = post(API + "/api/analysis/site", data=data, headers=cls.g_headers())
        if res is None: return '0.0.0.0'
        if res.json()['code'] == 401:
            TOKEN = cls.login()
            return cls.site_get_ip(url)
        try:
            return res.json()['data']['ip']
        except:
            return '0.0.0.0'

    # 获取站点的分析信息
    @classmethod
    def site_get_info(cls, url: str) -> dict:
        global TOKEN
        site = base64.b64encode(url.encode('utf-8')).decode()
        data = "type=basic&site={}".format(site)
        res = post(API + "/api/analysis/site", data=data, headers=cls.g_headers())
        if res is None: return {}
        if res.json()['code'] == 401:
            TOKEN = cls.login()
            return cls.site_get_info(url)
        try:
            return res.json()['data']
        except:
            return {}

    # 获取站点的分析信息
    @classmethod
    def site_get_cms(cls, url: str) -> dict:
        global TOKEN
        site = base64.b64encode(url.encode('utf-8')).decode()
        data = "type=cms&site={}".format(site)
        res = post(API + "/api/analysis/site", data=data, headers=cls.g_headers())
        if res is None: return {}
        if res.json()['code'] == 401:
            TOKEN = cls.login()
            return cls.site_get_info(url)
        try:
            return res.json()['data']
        except:
            return {}

    # 获取IP属地信息
    @classmethod
    def ip_get_info(cls, ip: str) -> dict:
        global TOKEN
        res = post(API + "/api/analysis/ip", data="type=shudi&ip=" + ip, headers=cls.g_headers(), timeout=12)
        if res is None: return {}
        if res.json()['code'] == 401:
            TOKEN = cls.login()
            return cls.ip_get_info(ip)
        try:
            logger.info((ip, res.json()['data']))
            return res.json()['data']
        except Exception as err:
            log.error(err)
            return {}

    # 检测URL的可用性
    @classmethod
    def url_check_status(cls, url: str, tim: int = 5, s: bool = False) -> Union[bool, Tuple[str, bool]]:
        log.info('check ' + url)
        try:
            get(url, timeout=tim)
            if s: return url, True
            return True
        except:
            if s: return url, False
            return False

    # 检测URL列表或单个URL的可用性
    @classmethod
    def urls_check_status(cls, urls: list or str, tim: int = 5) -> Dict[str, List[str]]:
        res = {
            "ok": [],
            "err": []
        }
        if type(urls) is str: urls = [urls]
        log.info(f'Count All urls: {len(urls)}')
        with cf.ThreadPoolExecutor(max_workers=10) as executor:
            checks = executor.map(lambda url: cls.url_check_status(url, tim, True), urls)

        for r in checks:
            res['ok'].append(r[0]) if r[1] else res['err'].append(r[0])
        return res


class Exploit:
    Info = None
    count = {'t': 0, "p": 0, 'v': 0}
    vul_info = list()

    @classmethod
    def run(cls, pl: bool = False):
        if cls.Info is None: cls.get_()
        t, p, v, i = '', '', '', ''
        while i != 'exit':
            res = cls.switch(t, p, v)
            if not t:
                i = input('[*] 请选择漏洞产品分类:')
                if i == '0' and not pl: cls.r_all(t, p, v)
                t = res[i] if i in res else ''
                continue

            elif not p:
                i = input('[*] 请选择产品(输入c返回上一层):')
                if i == '0' and not pl: cls.r_all(t, p, v)
                if i.startswith('c'):
                    t = ''
                    continue
                p = res[i] if i in res else ''
                continue

            elif not v:
                i = input('[*] 请选择漏洞(输入c返回上一层):')
                if i == '0' and not pl: cls.r_all(t, p, v)
                if i.startswith('c'):
                    p = ''
                    continue
                v = res[i] if i in res else ''
                if v:
                    if pl: return t, p, v
                    cls.r_all(t, p, v)
            else:
                t, p, v = '', '', ''
        sys.exit('退出')

    @classmethod
    def switch(cls, t: str = '', p: str = '', v: str = '') -> dict:
        tt, pp, vv = {}, {}, {}
        if not t:
            print("[+] 漏洞分类列表")
            print('', '>', '0、', '全部漏洞')
            for n, i in enumerate(sorted(list(cls.Info.keys()), reverse=True)):
                tt.update({f'{n + 1}': i})
                print('', '>', f'{n + 1}、', Model[i])
            return tt

        if not p and t in cls.Info:
            print("[+] 产品列表")
            print('', '>', '0、', '分类下全部漏洞')
            for n, i in enumerate(list(cls.Info[t].keys())):
                pp.update({f'{n + 1}': i})
                print('', '>', f'{n + 1}、', i)
            return pp

        if not v and t in cls.Info and p in cls.Info[t]:
            print("[+] 漏洞")
            print('', '>', '0、', '所有漏洞')
            for n, i in enumerate(list(cls.Info[t][p])):
                vv.update({f'{n + 1}': i})
                print('', '>', f'{n + 1}、', i)
            return vv
        return {}

    @classmethod
    def r_all(cls, t: str, p: str, v: str):
        cls.vul_info = list()
        url = input('请输入URL>>').strip()
        o_file = urlparse(url).netloc.split(':')[0] + Fun.timer()
        if t and t in cls.Info and p and p in cls.Info[t] and v and v in cls.Info[t][p]:
            poc = cls.get_poc(t, p, v)
            rs = cls.v_run(url, poc, o_file)
            if rs[0]: log.info(f"存在漏洞 {rs[2]['vName']}")
            if rs[0]: log.info('保存docx {}'.format(rs[1]))
            cls.vul_info.append({'url': url, 'file': rs[1], 'isVul': rs[0], 'pocRs': rs[2]})
        elif t and t in cls.Info and p and p in cls.Info[t]:
            for i in cls.Info[t][p]:
                poc = cls.get_poc(t, p, i)
                rs = cls.v_run(url, poc, o_file)
                if rs[0]: log.info('存在漏洞 {}'.format(rs[2]['vName']))
                if rs[0]: log.info('保存docx {}'.format(rs[1]))
                cls.vul_info.append({'url': url, 'file': rs[1], 'isVul': rs[0], 'pocRs': rs[2]})
        elif t and t in cls.Info:
            for m in cls.Info[t]:
                for _ in cls.Info[t][m]:
                    poc = cls.get_poc(t, m, _)
                    rs = cls.v_run(url, poc, o_file)
                    if rs[0]: log.info('存在漏洞 {}'.format(rs[2]['vName']))
                    if rs[0]: log.info('保存docx {}'.format(rs[1]))
                    cls.vul_info.append({'url': url, 'file': rs[1], 'isVul': rs[0], 'pocRs': rs[2]})
        else:
            for i in cls.Info:
                for o in cls.Info[i]:
                    for v in cls.Info[i][o]:
                        poc = cls.get_poc(i, o, v)
                        rs = cls.v_run(url, poc, o_file)
                        if rs[0]: log.info('存在漏洞 {}'.format(rs[2]['vName']))
                        if rs[0]: log.info('保存docx {}'.format(rs[1]))
                        cls.vul_info.append({'url': url, 'file': rs[1], 'isVul': rs[0], 'pocRs': rs[2]})

        y_vul = [i['file'] for i in cls.vul_info if i['isVul']]
        log.info('\r[+] 本次使用poc {} {}'.format(len(cls.vul_info), '个'))
        log.info('\r[+] 存在漏洞 {} {}'.format(len(y_vul), '个'))
        log.info("\r[*] 漏洞报告")
        for i in y_vul: print('>>', i)

    @classmethod
    def v_run(cls, url: str, poc: dict, o_file: str) -> tuple:
        log.info(f'\r{poc["vul_name"]}')
        if not os.path.exists(p_doc): os.mkdir(p_doc)
        run = cls.exploit(poc, url)
        if run['isVul']:
            if not os.path.exists(os.path.join(p_doc, o_file)): os.mkdir(os.path.join(p_doc, o_file))
            return True, Fun.write_docx(run, url, o_file), run
        return False, '', run

    @classmethod
    def init(cls):
        if cls.Info is None: cls.Info = cls.get_info()
        if not cls.Info: sys.exit('获取漏洞信息失败，程序退出')

    @classmethod
    def get_(cls, t: str = '') -> list:
        if cls.Info is None: cls.init()
        if t == 'model':
            return list(cls.Info.keys())
        else:
            cls.count['t'] = len(list(cls.Info.keys()))
            cls.count['p'], cls.count['v'] = 0, 0
            for p in cls.Info:
                cls.count['p'] += len(cls.Info[p].keys())
                for v in cls.Info[p]: cls.count['v'] += len(cls.Info[p][v])
        return cls.Info

    @classmethod
    def get_info(cls) -> dict:
        global TOKEN
        try:
            res = post(
                url=API + "/api/vul",
                data={'type': "_list", "model": "all"},
                headers={'token': TOKEN}
            )
            if res.json()['code'] == 401:
                TOKEN = secAPI.login()
                return cls.get_info()
            with open('config/config.json', 'r') as _f:
                _c = load(_f)
            _c.update({'vInfo': res.json()['data']})
            with open('config/config.json', 'w') as _f:
                dump(_c, _f)
            return res.json()['data']
        except Exception as err:
            log.warning(err)
            with open('config/config.json', 'r') as _f:
                _c = load(_f)
                return _c['vInfo'] if 'vInfo' in _c else {}

    @classmethod
    def check_pocs(cls):
        if cls.Info is None: cls.init()
        for i in cls.Info:
            for o in cls.Info[i]:
                for v in cls.Info[i][o]: cls.get_poc(i, o, v)

    @classmethod
    def get_poc(cls, t: str, p: str, v: str, enforce: bool = False) -> dict:
        log.info(f'POC {t, p, v}')
        global TOKEN
        if not os.path.exists(p_poc): os.mkdir(p_poc)
        file = os.path.join(p_poc, Fun.md5(Fun.get_computer_hash() + t + p + v))
        if os.path.exists(file) and enforce is False:
            with open(file) as _f:
                poc = load(_f)
                if poc: return poc
        log.info(f'Poc download => {t, p, v}')
        t, p, v = Fun.b64(t), Fun.b64(p), Fun.b64(v)
        res = post(API + "/api/vul", data={'type': "poc", "model": t, "product": p, "vul": v}, headers={'token': TOKEN})
        if res is None: return {}
        if res.json()['code'] == 401:
            TOKEN = secAPI.login()
            poc = cls.get_info()
        else:
            poc = res.json()['data']
        with open(file, 'w') as _f:
            dump(poc, _f)
        return poc

    @classmethod
    def exploit(cls, poc: dict, url: str) -> dict:
        result = {
            'isVul': False, "product": '', "vName": '', "vType": "", "vId": '', "level": '', "vDesc": '', "link": '',
            "version": "", "note": "", "search": {}, "fix": [],
            "request": {"url": "", "method": "", "header": {}, "body": ""},
            "response": {"code": 0, "url": "", "header": {}, "body": "", "time": 0}
        }
        timer = 0
        result.update({
            "product": poc['product'], "vName": poc['vul_name'], "vId": poc['vul_id'], "level": level[poc['level']],
            "version": poc['version'], "vDesc": poc['vul_desc'], "link": poc['link'], 'vType': poc['type'],
            "search": poc['search'], "fix": poc['fix'], "note": poc['note']
        })
        if 'model' in poc and poc['model']: return result
        if poc['protocol'] == 'http':
            exploits = poc['exploit'] if type(poc['exploit']) is list else [poc['exploit']]
            is_vul = None
            res = None
            _logic = poc['logic'] if 'logic' in poc and poc['logic'] else 'and'
            for exploit in exploits:
                exp = exploit['exp']
                _json = exp['json'] if 'json' in exp and exp['json'] else None
                file = exp['file'] if 'file' in exp and exp['file'] else None
                data = exp['data'] if 'data' in exp and exp['data'] else None
                if (_json or file or data) and exp['method'] not in ['put', 'post']: exp['method'] = 'post'
                _m = post if exp['method'] == 'post' else put if exp['method'] == 'put' else get
                link = url + '/' + exp['path'].lstrip('/')
                # 漏洞访问URL
                header = exp['header'] if 'header' in exp else {}
                header.update({"Cookie": exp['cookie']}) if 'cookie' in exp and exp['cookie'] else None
                header = g_headers(link, header)
                # 漏洞指定Header头
                timer = time()
                try:
                    if exp['method'] in ['put']:
                        res1 = _m(link, headers=header, data=exp['data'], allow_redirects=exp['follow'], verify=False,
                                  timeout=Timeout)
                    elif exp['method'] in ['post'] or (_json or file or data):
                        res1 = _m(link, headers=header, json=_json, files=file, data=data, verify=False,
                                  timeout=Timeout, allow_redirects=exp['follow'])
                    else:
                        res1 = _m(link, headers=header, verify=False, timeout=Timeout, allow_redirects=exp['follow'])
                except:
                    return result
                timer = time() - timer
                if res1 is None: return result
                checks = exploit['verify']['check']
                # 漏洞尝试方法列表
                logic = exploit['verify']['logic']
                # 当有多个判断条件的时候 -> 判断逻辑
                is_vul = True if not checks else False
                # 当逻辑为and的时候，默认为True，否则False
                res = res1
                if 'path' in exploit['verify'] and exploit['verify']['path']:
                    timer = time()
                    ver_link = url + '/' + exploit['verify']['path'].lstrip('/')
                    try:
                        res = get(ver_link, headers=g_headers(ver_link))
                    except:
                        return result
                    timer = timer - time()
                if res is None: return result
                for check in checks:
                    err = True
                    # 当结果有一次为False的时候，err为False且以后不发生改变
                    if 'status' in check and err: err = res.status_code == check['status']
                    if 'body' in check and err: err = bool(re.search(re.escape(check['body']), res.text))
                    if 'nbody' in check and err: err = not bool(re.search(re.escape(check['nbody']), res.text))
                    if "header" in check and check['header'] and err:
                        key = list(check['header'].keys())[0]
                        err = key in res.headers and (
                                not check['header'][key] or re.search(check['header'][key], res.headers[key]))
                    if "length" in check and err:
                        length = check['length']
                        if length['check'] == '>=':
                            err = length(res.text) >= length['num']
                        elif length['check'] == '<=':
                            err = length(res.text) <= length['num']
                        elif length['check'] == '==':
                            err = length(res.text) == length['num']
                        elif length['check'] == '!=':
                            err = length(res.text) != length['num']
                    if "time" in check and err:
                        c_time = check['time']
                        if c_time['check'] == '>=':
                            err = timer >= c_time['num']
                        elif c_time['check'] == '<=':
                            err = timer <= c_time['num']
                        elif c_time['check'] == '==':
                            err = timer == c_time['num']
                        elif c_time['check'] == '!=':
                            err = timer != c_time['num']
                    is_vul = err
                    if logic == 'or':
                        if is_vul is False: continue
                        if is_vul: break
                    if logic == 'and':
                        if is_vul: continue
                        if is_vul is False: break
                if not result['request']['url']:
                    result['request']['url'] = res1.request.url
                    result['request']['header'].update(res1.request.headers)
                    result['request']['method'] = res1.request.method
                    try:
                        result['request']['body'] = res1.request.body.decode() if res1.request.body else ''
                    except:
                        result['request']['body'] = res1.request.body if res1.request.body else ''
                if (_logic == 'and' and not is_vul) or (_logic == 'or' and is_vul): break

            result['isVul'] = is_vul
            result['response']['code'] = res.status_code
            result['response']['time'] = round(timer, 2)
            result['response']['url'] = res.url
            result['response']['header'].update(res.headers)
            try:
                result['response']['body'] = res.text
            except:
                result['response']['body'] = ''
        return result


if __name__ == '__main__':
    e = Exploit()
    e.get_()
    e.check_pocs()
    log.info('漏洞分类: {} {}'.format(e.count['t'], '个'))
    log.info('漏洞产品: {} {}'.format(e.count['p'], '个'))
    log.info('漏洞Poc: {} {}'.format(e.count['v'], '个'))
    print()
    e.run()
