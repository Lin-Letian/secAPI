from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from Script.request import async_get, async_post, async_put, r_post, r_get, r_put
from Script.function import Function
from Script import Document, time, sleep
from urllib.parse import urlparse
from Script.engine import Engine
from docx.shared import RGBColor
from Script.action import Action
from json import load, loads
import os, re

level = {0: "无危害", 1: '低危', 2: '中危', 3: '高危'}
Model = {
    "auth": "统一/认证系统",
    "cms": "CMS",
    "Components": "组件",
    "Database": "数据库",
    "Framework": "框架",
    "NetDrivers": "网络设备",
    "OA": "办公系统",
    "ServicesApplication": "服务器应用",
    "SecurityApp": "安全产品",
    "VideoPanel": "视频监控平台",
    "WebApplication": "Web应用程序/系统",
    "WebServer": "Web中间件/服务器"
}
vPath = os.path.join(Function.path, 'Vuls')


# 获取dnslog.cn结果
async def dnslog_get_domain():
    result = {'dom': '', 'cookie': ''}
    try:
        res = await async_get('http://dnslog.cn/getdomain.php', timeout=10)
        text = await res.text()
        result.update({'dom': text, "cookie": f"PHPSESSID={res.cookies.get('PHPSESSID')}"})
    except:
        pass
    return result


# 获取dnslog.cn结果
async def dnslog_get_log(cookie: str):
    header = {'Cookie': cookie, "Referer": "http://dnslog.cn/"}
    for i in ['', '']:
        try:
            res = await async_get('http://dnslog.cn/getrecords.php', timeout=10, header=header)
            loads(await res.text())
            if res: return True
        except:
            pass
        sleep(.5)
    return False


async def get_list(model, p_type: str = None, product: str = None):
    vul = dict()
    _m = [path for path in os.listdir(vPath) if os.path.isdir(os.path.join(vPath, path))]
    models = sorted([[i, Model[i]] for i in _m if i in Model], reverse=True)
    model = model.strip() if model is not None else model
    p_type = p_type.strip() if p_type is not None else p_type
    product = product.strip() if product is not None else product
    if model == 'model': return models
    for v in models:
        m = v[0]
        if m not in vul: vul.update({m: {}})
        p_path = os.path.join(vPath, m)
        products = [path for path in os.listdir(p_path) if os.path.isdir(os.path.join(p_path, path))]
        for p in products:
            if m not in vul[m]: vul[m].update({p: []})
            pv_path = os.path.join(p_path, p)
            vul_ = [os.path.splitext(item)[0] for item in os.listdir(pv_path) if
                    os.path.isfile(os.path.join(pv_path, item))]
            vul[m][p].extend(vul_)

    if model == 'all': return vul

    if model == 'select' and p_type is not None:
        if p_type in vul and product is not None and product in vul[p_type]: return sorted(vul[p_type][product],
                                                                                           reverse=True)
        if p_type in vul: return sorted(list(vul[p_type].keys()), reverse=True)
    return []


async def get_poc(data: list or tuple) -> dict:
    m, p, v = data
    v_file = os.path.join(vPath, m, p, v + '.json')
    if not os.path.exists(v_file): return {}
    with open(v_file, encoding='utf-8') as f: return load(f)


async def r_exploit(data: list) -> dict:
    url, m, p, v = data
    result = {
        'isVul': False, "product": '', "vName": '', "vType": "", "vId": '', "level": '', "vDesc": '', "link": '',
        "version": "", "note": "", "search": {}, "fix": [],
        "request": {"url": "", "method": "", "header": {}, "body": ""},
        "response": {"code": 0, "url": "", "header": {}, "body": "", "time": 0}
    }
    poc = await get_poc([m, p, v])
    timer = 0
    if not poc: return result
    result.update({
        "product": poc['product'], "vName": poc['vul_name'], "vId": poc['vul_id'], "level": level[poc['level']],
        "version": poc['version'], "vDesc": poc['vul_desc'], "link": poc['link'], 'vType': poc['type'],
        "search": poc['search'], "fix": poc['fix'], "note": poc['note']
    })

    if poc['protocol'] == 'http':
        exploits = poc['exploit'] if type(poc['exploit']) is list else [poc['exploit']]
        is_vul, res = None, None
        _logic = poc['logic'] if 'logic' in poc and poc['logic'] else 'and'
        for _exp in exploits:
            _mod = dict()
            if 'model' in poc and poc['model']:
                if poc['model'] == 'dnslog':
                    _mod = await dnslog_get_domain()
                    if not _mod['dom'] or not _mod['cookie']: return result

            exp = _exp['exp']
            _json = exp['json'] if 'json' in exp and exp['json'] else None
            file = exp['file'] if 'file' in exp and exp['file'] else None
            data = exp['data'] if 'data' in exp and exp['data'] else None
            if 'model' in poc and poc['model'] == 'dnslog':
                data = data.replace('{{m:rep}}', _mod['dom']) if type(data) is str else data
                exp['path'] = exp['path'].replace('{{m:rep}}', _mod['dom']) if exp['path'] else exp['path']
            if (_json or file or data) and exp['method'] not in ['put', 'post']: exp['method'] = 'post'
            _method = r_post if exp['method'] == 'post' else r_put if exp['method'] == 'put' else r_get
            link = url + '/' + exp['path'].lstrip('/')
            # 漏洞访问URL
            header = exp['header'] if 'header' in exp else {}
            header.update({"Cookie": exp['cookie']}) if 'cookie' in exp and exp['cookie'] else None
            # 漏洞指定Header头
            timer = time()
            if exp['method'] in ['put']:
                res1 = _method(link, header=header, data=exp['data'], allow_redirects=exp['follow'])
            elif exp['method'] in ['post'] or (_json or file or data):
                res1 = _method(link, header=header, json=_json, files=file, data=data)
            else:
                res1 = _method(link, header=header)
            timer = time() - timer
            if res1 is None: return result
            checks = _exp['verify']['check']
            # 漏洞尝试方法列表
            logic = _exp['verify']['logic']
            # 当有多个判断条件的时候 -> 判断逻辑
            is_vul = True if not checks else False
            # 当逻辑为and的时候，默认为True，否则False
            res = res1
            if 'path' in _exp['verify'] and _exp['verify']['path']:
                timer = time()
                res = r_get(url + '/' + _exp['verify']['path'].lstrip('/'))
                timer = timer - time()
            if res is None: return result
            if not result['request']['url']:
                result['request']['url'] = str(res1.request.url)
                result['request']['header'].update(res1.request.headers)
                result['request']['method'] = res1.request.method
                try:
                    result['request']['body'] = res1.request.body.decode()
                except:
                    result['request']['body'] = str(res1.request.body)
            for check in checks:
                err = True
                text = Function.beautifulsoup(res.text, res.content, dict(res.headers)).decode()
                if Function.is_waf(text): return result
                # 当结果有一次为False的时候，err为False且以后不发生改变
                if 'status' in check and err: err = res.status_code == check['status']
                if 'body' in check and err: err = bool(re.search(re.escape(check['body']), text))
                if 'nbody' in check and err: err = not bool(re.search(re.escape(check['nbody']), text))
                if "header" in check and check['header'] and err:
                    key = list(check['header'].keys())[0]
                    err = key in dict(res.headers) and (
                            not check['header'][key] or re.search(check['header'][key], res.headers[key]))
                if "length" in check and err:
                    length = check['length']
                    if length['check'] == '>=':
                        err = length(res.text) >= length['num']
                    elif length['check'] == '<=':
                        err = length(res.text) <= length['num']
                    elif length['check'] == '==':
                        err = length(res.text) == length['num']
                    elif length['check'] == '!=':
                        err = length(res.text) != length['num']
                if "time" in check and err:
                    c_time = check['time']
                    if c_time['check'] == '>=':
                        err = timer >= c_time['num']
                    elif c_time['check'] == '<=':
                        err = timer <= c_time['num']
                    elif c_time['check'] == '==':
                        err = timer == c_time['num']
                    elif c_time['check'] == '!=':
                        err = timer != c_time['num']
                if 'model' in check:
                    if check['model'] == 'dnslog':
                        err = await dnslog_get_log(_mod['cookie'])
                is_vul = err
                if logic == 'or':
                    if is_vul is False: continue
                    if is_vul: break
                if logic == 'and':
                    if is_vul: continue
                    if is_vul is False: break
            if (_logic == 'and' and not is_vul) or (_logic == 'or' and is_vul): break

        result['isVul'] = is_vul
        result['response']['code'] = res.status_code
        result['response']['time'] = round(timer, 2)
        result['response']['url'] = str(res.url)
        result['response']['header'].update(res.headers)
        try:
            result['response']['body'] = Function.beautifulsoup(res.text, res.content, dict(res.headers)).decode()
        except:
            result['response']['body'] = res.text
    if result['isVul']: result['file'] = await write_docx(result, url)
    return result


async def _w_icp(doc, text: str) -> dict:
    dom = urlparse(text).netloc.split(':')[0]
    if dom[0] in ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9']: return doc
    icp = await Engine.domain_icp_query(dom)
    if icp['size'] > 0:
        doc.add_paragraph()
        doc.add_paragraph()
        heading = doc.add_heading('ICP备案查询信息', 1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        icp = icp['list'][0]
        _row_tit(doc, "单位名称").add_run(icp['unitName'])
        _row_tit(doc, "单位性质").add_run(icp['natureName'])
        _row_tit(doc, "ICP号").add_run(icp['serviceLicence'])
        _row_tit(doc, "域名").add_run(icp['domain'])
        _row_tit(doc, "更新时间").add_run(icp['updateRecordTime'])
    return doc


def _row_tit(doc, text: str):
    line = doc.add_paragraph()
    tit = line.add_run(f'{text}：')
    tit.bold = True
    tit.font.color.rgb = RGBColor(0, 51, 102)  # 使用RGB颜色
    return line


async def write_docx(v: dict, url: str) -> str:
    fname = Function.md5(f'{v["vName"]}{time()}{v["level"]}')
    doc = Document()
    # 标题
    heading = doc.add_heading('漏洞报告', 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    t = doc.add_paragraph('生成时间：' + Action.timer())
    t.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    _row_tit(doc, "漏洞名称").add_run(v['vName'])
    _row_tit(doc, "漏洞URL").add_run(url)
    _row_tit(doc, "是否存在漏洞").add_run('是' if v['request'] else '否')
    _row_tit(doc, "影响产品").add_run(v['product'])
    _row_tit(doc, "影响版本").add_run(v['version'])
    _row_tit(doc, "危险等级").add_run(v['level'])
    if v['vId']:
        _row_tit(doc, "漏洞编号")
        for n, i in enumerate(v['vId']):
            doc.add_paragraph(str(n + 1) + '：' + i)

    _row_tit(doc, "漏洞描述").add_run(v['vDesc'])
    _row_tit(doc, "参考链接").add_run(v['link'])
    if v['fix']:
        _row_tit(doc, "修复建议")
        for n, i in enumerate(v['fix']):
            doc.add_paragraph(str(n + 1) + '：' + i)
    doc.add_paragraph('')
    doc.add_paragraph('')
    request = v['request']
    _row_tit(doc, ">> 请求")
    _row_tit(doc, "URL").add_run(request['url'])
    _row_tit(doc, "请求方法").add_run(request['method'])
    _row_tit(doc, "请求头")
    table = doc.add_table(rows=1, cols=1)
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            hed, text = request['header'], ''
            for i in hed: text += i + ': ' + hed[i] + '\n'
            cell.text = text.rstrip('\n')
    if request['body']:
        _row_tit(doc, "请求体")
        table = doc.add_table(rows=1, cols=1)
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells): cell.text = request['body']
    doc.add_paragraph('')
    doc.add_paragraph('')
    response = v['response']
    _row_tit(doc, ">> 响应")
    _row_tit(doc, "URL").add_run(response['url'])
    _row_tit(doc, "状态码").add_run(str(response['code']))
    _row_tit(doc, "响应头")
    table = doc.add_table(rows=1, cols=1)
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            hed, text = response['header'], ''
            for i in hed: text += i + ': ' + hed[i] + '\n'
            cell.text = text.rstrip('\n')
    _row_tit(doc, "响应体")
    table = doc.add_table(rows=1, cols=1)
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells): cell.text = response['body']

    await _w_icp(doc, url)
    doc.save(os.path.join(os.path.join(Function.path, 'files'), fname + '.docx'))
    return fname
